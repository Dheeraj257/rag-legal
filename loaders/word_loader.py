from docx import Document
from langchain_core.documents import Document as LCDocument
from pathlib import Path

def load_doc(file_path):

    doc = Document(file_path)
    word_doc = []
    file_name = Path(file_path).name

    current_para = "Document starts here"
    current_heading = None
    para_buffer = []
    para_index = 0

    for element in doc.element.body:
        tag = element.tag.split("}")[-1]

        if tag == "p":
            if para_index >= len(doc.paragraphs):
                continue
            para = doc.paragraphs[para_index]
            para_index += 1

            text = para.text.strip()
            style = para.style.name

            if not text:
                continue

            if style =="Heading 1":
                if para_buffer:
                    word_doc.append(LCDocument(
                        page_content="\n".join(para_buffer),
                        metadata={
                            "source": file_name,
                            "para": current_para,
                            "heading": current_heading,
                            "citation": f"{file_name}: {current_para}" 
                        }
                    ))
                current_heading = text
                current_para = text
                para_buffer = []
            
            elif style == "Heading 2":
                if para_buffer:
                    word_doc.append(LCDocument(
                        page_content= "\n".join(para_buffer),
                        metadata={
                            "source": file_name,
                            "para": current_para,
                            "heading": current_heading,
                            "citation": f"{file_name}: {current_para}"
                        }
                    ))
                current_heading = text
                current_para = text
                para_buffer = []
            else:
                para_buffer.append(text)
    
        elif tag == "tbl":
            if para_buffer:
                    word_doc.append(LCDocument(
                        page_content= "\n".join(para_buffer),
                        metadata={
                            "source": file_name,
                            "para": current_para,
                            "heading": current_heading,
                            "citation": f"{file_name}: {current_para}"
                        }
                    ))
            para_buffer = []

            for table in doc.tables:
                if table._element is element:
                    rows = []
                    for row in table.rows:
                        cleaned = [cell.text.strip() for cell in row.cells]
                        rows.append(" | ".join(cleaned))
                    table_content = "\n".join(rows)

                    word_doc.append(LCDocument(
                        page_content= f"[TABLE] \n\n{table_content}",
                        metadata={
                            "source":file_name,
                        }
                    ))
    if para_buffer:
        word_doc.append(LCDocument(
                        page_content= "\n".join(para_buffer),
                        metadata={
                            "source": file_name,
                            "para": current_para,
                            "heading": current_heading,
                            "citation": f"{file_name}: {current_para}"
                        }
                    ))
    return word_doc


            

        
